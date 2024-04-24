import os
from docx import Document

# Function to combine all .docx files in a folder
def combine_docx(folder_path, combined_filename):
    combined_doc = Document()
    
    # Get a list of all .docx filenames in the folder and sort them based on chapter numbers
    filenames = sorted(os.listdir(folder_path), key=lambda x: int(x.split("_")[1].split(".")[0]))
    
    for filename in filenames:
        if filename.endswith(".docx"):
            doc_path = os.path.join(folder_path, filename)
            doc = Document(doc_path)
            
            # Add chapter heading
            chapter_heading = doc.paragraphs[0].text.strip()  # Extract the first line of the chapter
            combined_doc.add_heading(chapter_heading, level=1)
            
            # Add content of the chapter
            for paragraph in doc.paragraphs[1:]:
                combined_doc.add_paragraph(paragraph.text)  # Add content of the chapter
            combined_doc.add_page_break()  # Add page break between chapters
    
    combined_doc.save(combined_filename)

# Function to count words in a document
def count_words(docx_path):
    doc = Document(docx_path)
    total_words = 0
    for paragraph in doc.paragraphs:
        total_words += len(paragraph.text.split())
    return total_words

# Function to export word count to a text file
def export_word_count(count, output_filename):
    with open(output_filename, "w") as file:
        file.write(str(count))

# Folder paths
input_folder = "Mutation of the Apocalypse"
output_folder = "Statistics"  # Adjusted output folder name

# Create output folder if it doesn't exist
if not os.path.exists(output_folder):
    os.makedirs(output_folder)

# Combine .docx files
combined_filename = os.path.join(output_folder, "Mutation_of_the_Apocalypse.docx")
if os.path.exists(combined_filename):
    os.remove(combined_filename)  # Remove existing file to replace it
combine_docx(input_folder, combined_filename)

# Count words
word_count = count_words(combined_filename)

# Export word count
count_filename = os.path.join(output_folder, "count.txt")
export_word_count(word_count, count_filename)

print("Combined document and word count exported successfully.")
