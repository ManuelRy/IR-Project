import os
import requests
from bs4 import BeautifulSoup
from docx import Document

# URL of the chapter
url = "https://www.royalroad.com/fiction/57452/mutation-of-the-apocalypse/chapter/1396465/chapter-105-silent-farewell"

# Fetch the HTML content of the webpage
response = requests.get(url)
soup = BeautifulSoup(response.text, 'html.parser')

# Find the element containing the chapter content
chapter_content = soup.find('div', class_='chapter-content')

# Extract the text content of the chapter
chapter_text = chapter_content.get_text(separator='\n')

# Extract the title of the chapter
title_element = soup.find('h1', class_='font-white')
chapter_title = title_element.text.strip() if title_element else "Untitled Chapter"

# Replace invalid characters in the title for file name
file_name = f"{chapter_title.replace(' ', '_').replace(':', '')}.docx"

# Create the directory if it doesn't exist
directory = "./Mutation of the Apocalypse"
if not os.path.exists(directory):
    os.makedirs(directory)

# Create a new document
doc = Document()

# Add the chapter content to the document
doc.add_heading(chapter_title, level=1)
doc.add_paragraph(chapter_text)

# Save the document with the title of the chapter
file_path = os.path.join(directory, file_name)
doc.save(file_path)

print(f"Chapter saved successfully at: {file_path}")
